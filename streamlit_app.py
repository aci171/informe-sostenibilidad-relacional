import streamlit as st
import os
import tempfile
import base64
from google import genai
from docx import Document
import fitz  # PyMuPDF para PDF

# Configuración de API
genai.configure(api_key=os.environ.get("GEMINI_API_KEY"))
model = genai.GenerativeModel("gemini-2.5-pro")

# Título de la App
st.set_page_config(page_title="Informe Relacional", layout="wide")
st.title("🧭 Generador de Informes con Sostenibilidad Relacional")

# Subida de documento o entrada manual
opcion = st.radio("¿Cómo quieres generar el informe?", ("Subir un documento", "Escribir manualmente"))

texto_fuente = ""

if opcion == "Subir un documento":
    archivo = st.file_uploader("Selecciona un archivo (.pdf, .docx, .txt)", type=["pdf", "docx", "txt"])
    if archivo:
        ext = archivo.name.split(".")[-1].lower()
        if ext == "txt":
            texto_fuente = archivo.read().decode("utf-8")
        elif ext == "docx":
            doc = Document(archivo)
            texto_fuente = "\n".join([p.text for p in doc.paragraphs])
        elif ext == "pdf":
            pdf_doc = fitz.open(stream=archivo.read(), filetype="pdf")
            texto_fuente = "\n".join(page.get_text() for page in pdf_doc)
else:
    texto_fuente = st.text_area("Escribe el texto base para el informe", height=300)

if texto_fuente:
    if st.button("🔍 Generar Informe"):
        with st.spinner("Generando informe..."):
            prompt = f"""Eres un redactor especializado en sostenibilidad relacional y cultura organizacional. 
Redacta un informe claro, cercano y bien estructurado basado en este contenido:\n\n{texto_fuente}"""
            respuesta = model.generate_content(prompt)
            informe = respuesta.text

            st.subheader("📄 Informe Generado")
            st.write(informe)

            # Guardar informe como .docx
            docx_file = Document()
            for par in informe.split("\n"):
                docx_file.add_paragraph(par)
            tmp_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
            docx_file.save(tmp_docx.name)

            # Mostrar botón para descargar
            with open(tmp_docx.name, "rb") as f:
                st.download_button("⬇️ Descargar Informe (.docx)", f, file_name="informe_relacional.docx")
